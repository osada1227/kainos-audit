"""再下請負通知書 審査 + 対話 Web フロント (Streamlit).

既存 CLI (chat_subcontract.py / audit_subcontract.py) と同じ system_prompt /
ロジックを再利用しつつ、画面で完結する。

起動:
    streamlit run app.py
"""
from __future__ import annotations

import base64
import io
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

import fitz  # PyMuPDF
import streamlit as st
import streamlit_authenticator as stauth
import yaml
from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).parent
sys.path.insert(0, str(ROOT / "system_prompt"))
from build_system_prompt import build_system_blocks  # noqa: E402

import anthropic  # noqa: E402


SESSIONS_DIR = ROOT / "sessions"
SESSIONS_DIR.mkdir(exist_ok=True)
SOURCES_DIR = ROOT / "sources"

IMAGE_MEDIA = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
}

# 構造化審査プロンプト (audit_subcontract.py と同等)
AUDIT_PROMPT = (
    "添付の再下請負通知書（全建統一様式第1号-甲）を、システムプロンプトの "
    "Step 1〜4 と <legal_corpus> に照らして厳格に審査し、指定の出力"
    "フォーマット（全体判定 / 法令違反 / 記載漏れ / 正常項目 / 読み取り"
    "不能 / 実務アドバイス）で Markdown 形式の審査結果を返してください。"
)

# 引用パターン: [ガイドライン第12版 p.15] / [施工体制台帳マニュアル p.3-4] 等
CITATION_RE = re.compile(
    r"\[([^\[\]]+?)\s*p\.\s*([0-9]+(?:-[0-9]+)?)\]"
)

# 引用元ドキュメント表記 → sources/ 内 PDF のマッピング
DOC_MAP: list[tuple[str, str]] = [
    ("ガイドライン", "guideline_v12_R8-1.pdf"),
    ("建設業法令遵守", "guideline_v12_R8-1.pdf"),
    ("チェックポイント", "sekoutaisei_checkpoint.pdf"),
    ("施工体制台帳等活用マニュアル", "sekoutaisei_manual.pdf"),
    ("施工体制台帳マニュアル", "sekoutaisei_manual.pdf"),
    ("関東地方整備局", "ktr_kensetsugyohou_R8-2.pdf"),
    ("建設業法", "ktr_kensetsugyohou_R8-2.pdf"),
    ("近畿地整", "kkr_sekoutaisei_270401.pdf"),
    ("近畿地方整備局", "kkr_all-data_R7-4.pdf"),
    ("建設経済", "kensetsu_keizai_001855436.pdf"),
    ("建設業課", "kensetsu_keizai_001855436.pdf"),
]


def resolve_source(doc_title: str) -> Path | None:
    for needle, filename in DOC_MAP:
        if needle in doc_title:
            p = SOURCES_DIR / filename
            if p.exists():
                return p
    return None


def highlight_citations(text: str) -> str:
    """[...p.N] を <mark> でハイライト表示する HTML 化."""
    return CITATION_RE.sub(
        r'<mark style="background:#fef3c7;padding:0 3px;'
        r'border-radius:3px;font-weight:600;">[\1 p.\2]</mark>',
        text,
    )


def extract_citations(text: str) -> list[dict]:
    seen = set()
    result = []
    for m in CITATION_RE.finditer(text):
        doc = m.group(1).strip()
        page = m.group(2).strip()
        key = (doc, page)
        if key in seen:
            continue
        seen.add(key)
        result.append({"doc": doc, "page": page, "source": resolve_source(doc)})
    return result


@st.cache_data(show_spinner=False, max_entries=128)
def render_pdf_page(pdf_path_str: str, page_num: int, zoom: float = 1.6) -> bytes | None:
    """1始まりのページ番号で PDF を PNG bytes 化。範囲外 / 失敗時は None."""
    pdf_path = Path(pdf_path_str)
    if not pdf_path.exists():
        return None
    try:
        doc = fitz.open(pdf_path)
    except Exception:
        return None
    try:
        idx = page_num - 1
        if idx < 0 or idx >= len(doc):
            return None
        page = doc.load_page(idx)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        return pix.tobytes("png")
    except Exception:
        return None
    finally:
        doc.close()


@st.cache_data(show_spinner=False, max_entries=32)
def get_pdf_page_count(pdf_path_str: str) -> int:
    pdf_path = Path(pdf_path_str)
    if not pdf_path.exists():
        return 0
    try:
        doc = fitz.open(pdf_path)
        n = len(doc)
        doc.close()
        return n
    except Exception:
        return 0


def _parse_first_page(page_str: str) -> int | None:
    head = page_str.split("-")[0].strip()
    try:
        return int(head)
    except ValueError:
        return None


def render_citations(citations: list[dict], key_prefix: str) -> None:
    """引用一覧を expander で描画。各引用に PDF ページ画像とダウンロードを付ける."""
    if not citations:
        return
    with st.expander(f"📚 引用 {len(citations)}件 (該当ページ画像つき)", expanded=False):
        for i, c in enumerate(citations):
            header = f"**{c['doc']}** p.{c['page']}"
            if c["source"]:
                header += f"  —  `sources/{c['source'].name}`"
            st.markdown(header)

            if not c["source"]:
                st.caption("⚠ 対応する原本PDFを sources/ から特定できませんでした。")
                st.divider()
                continue

            first_page = _parse_first_page(c["page"])
            if first_page is None:
                st.caption(f"⚠ ページ番号 '{c['page']}' を解析できませんでした。")
            else:
                page_count = get_pdf_page_count(str(c["source"]))
                if page_count <= 0:
                    st.caption(f"⚠ PDFを開けませんでした ({c['source'].name}).")
                else:
                    state_key = f"pgnav_{key_prefix}_{i}"
                    if state_key not in st.session_state:
                        st.session_state[state_key] = min(
                            max(first_page, 1), page_count
                        )
                    current = st.session_state[state_key]
                    nav1, nav2, nav3, nav4 = st.columns([1, 1, 3, 2])
                    with nav1:
                        prev_clicked = st.button(
                            "◀ 前",
                            key=f"prev_{key_prefix}_{i}",
                            disabled=current <= 1,
                            use_container_width=True,
                        )
                    with nav2:
                        next_clicked = st.button(
                            "次 ▶",
                            key=f"next_{key_prefix}_{i}",
                            disabled=current >= page_count,
                            use_container_width=True,
                        )
                    with nav4:
                        reset_clicked = st.button(
                            f"↺ 引用 p.{first_page}",
                            key=f"reset_{key_prefix}_{i}",
                            disabled=current == first_page,
                            use_container_width=True,
                        )
                    if prev_clicked:
                        current = max(1, current - 1)
                    if next_clicked:
                        current = min(page_count, current + 1)
                    if reset_clicked:
                        current = first_page
                    current = max(1, min(current, page_count))
                    st.session_state[state_key] = current
                    with nav3:
                        st.caption(
                            f"ページ **{current}** / {page_count}  "
                            f"（引用は p.{first_page}）"
                        )

                    png = render_pdf_page(str(c["source"]), current)
                    if png:
                        st.image(
                            png,
                            caption=f"{c['source'].name} — p.{current}",
                            use_container_width=True,
                        )
                        if current == first_page:
                            st.caption(
                                "※ 引用ページ番号は AI 推定値。前後ページもご確認を。"
                            )
                    else:
                        st.caption(
                            f"⚠ p.{current} の描画に失敗 ({c['source'].name})."
                        )

            try:
                pdf_bytes = c["source"].read_bytes()
                st.download_button(
                    label=f"📥 {c['source'].name} をダウンロード",
                    data=pdf_bytes,
                    file_name=c["source"].name,
                    mime="application/pdf",
                    key=f"dl_{key_prefix}_{i}",
                    use_container_width=False,
                )
            except OSError:
                pass

            st.divider()


# ─────────────────────────────────────────────────────────────────────
#  審査結果の構造化サマリー（Markdown heuristic parser）
# ─────────────────────────────────────────────────────────────────────

# (見出しキーワード, 絵文字, セクションキー, 早期展開するか)
SECTION_PATTERNS: list[tuple[str, str, str, bool]] = [
    ("全体判定", "📄", "overall", False),
    ("法令違反", "🚨", "critical", True),
    ("記載漏れ", "⚠", "warning", False),
    ("正常に確認", "✅", "ok", False),
    ("読み取り不能", "🔍", "unclear", False),
    ("最新法改正", "💡", "advice", False),
    ("実務アドバイス", "💡", "advice", False),
]


def parse_audit_sections(md: str) -> dict[str, list[str]]:
    """審査Markdownを {section_key: [bullet_items]} に分解."""
    sections: dict[str, list[str]] = {}
    current_key: str | None = None
    current_items: list[str] = []
    heading_re = re.compile(r"^#{2,4}\s+(.*?)\s*$")
    bullet_re = re.compile(r"^\s*[*\-+]\s+(.*)$")

    for line in md.splitlines():
        m_h = heading_re.match(line)
        if m_h:
            if current_key is not None:
                sections[current_key] = sections.get(current_key, []) + current_items
            current_items = []
            heading_text = m_h.group(1)
            current_key = None
            for needle, _emoji, key, _exp in SECTION_PATTERNS:
                if needle in heading_text:
                    current_key = key
                    break
            continue
        if current_key is None:
            continue
        m_b = bullet_re.match(line)
        if m_b:
            current_items.append(m_b.group(1).strip())
        elif current_key == "overall":
            # 全体判定は段落形式なので非空行を拾う（例示の括弧書きは除外）
            stripped = line.strip()
            if stripped and not stripped.startswith("（例") and not stripped.startswith("(例"):
                current_items.append(stripped)
    if current_key is not None:
        sections[current_key] = sections.get(current_key, []) + current_items
    return sections


def render_checklist_summary(md: str, position: str = "above") -> None:
    """KPIカード + セクションexpanderで審査結果を要約描画.

    position: "above" (live response 後など、本文と一緒に下に置く) / "history"
    """
    sections = parse_audit_sections(md)
    has_any = any(sections.get(k) for k in ("critical", "warning", "ok", "unclear"))
    if not has_any:
        # AUDIT 形式でない応答（自由Q&A等）はサマリーを出さない
        return

    with st.container(border=True):
        label = (
            "### 📋 審査サマリー (履歴)"
            if position == "history"
            else "### 📋 審査サマリー"
        )
        st.markdown(label)
        cols = st.columns(4)
        kpis = [
            (cols[0], "🚨 重大エラー", "critical"),
            (cols[1], "⚠ 記載漏れ", "warning"),
            (cols[2], "✅ 正常項目", "ok"),
            (cols[3], "🔍 要再確認", "unclear"),
        ]
        for col, lab, key in kpis:
            col.metric(lab, len(sections.get(key, [])))

        # 全体判定があれば最上段に
        overall = sections.get("overall", [])
        if overall:
            st.markdown("**📄 全体判定**")
            for item in overall:
                st.markdown(
                    f"- {highlight_citations(item)}",
                    unsafe_allow_html=True,
                )

        # 各セクションの詳細
        for needle, emoji, key, expanded_default in SECTION_PATTERNS:
            if key in ("overall",):
                continue
            items = sections.get(key, [])
            if not items:
                continue
            seen_label = f"{emoji} {needle} ({len(items)}件)"
            with st.expander(seen_label, expanded=expanded_default):
                for item in items:
                    st.markdown(
                        f"- {highlight_citations(item)}",
                        unsafe_allow_html=True,
                    )


# ─────────────────────────────────────────────────────────────────────
#  Markdown → docx エクスポート
# ─────────────────────────────────────────────────────────────────────

# インライン記法: **bold** / `code` / [doc p.N] 引用
INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*|`[^`]+?`|\[[^\[\]]+?\s*p\.[0-9]+(?:-[0-9]+)?\])"
)


def _add_runs(paragraph, text: str) -> None:
    """インライン要素を paragraph の run に分解して追加."""
    if not text:
        return
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x6B, 0x21, 0xA8)
        elif part.startswith("[") and part.endswith("]") and "p." in part:
            run = paragraph.add_run(part)
            run.bold = True
            run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
        else:
            paragraph.add_run(part)


@st.cache_data(show_spinner=False, max_entries=64)
def convert_markdown_to_docx(md: str, title: str = "再下請負通知書 審査結果") -> bytes:
    """Markdown文字列 (主に審査出力) を docx の bytes に変換."""
    doc = Document()
    doc.add_heading(title, level=0)

    ts_line = doc.add_paragraph()
    run = ts_line.add_run(
        f"生成日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    heading_re = re.compile(r"^(#{1,4})\s+(.*)$")
    bullet_re = re.compile(r"^(\s*)([*\-+])\s+(.*)$")

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        m_h = heading_re.match(line)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            doc.add_heading(m_h.group(2).strip(), level=level)
            continue
        m_b = bullet_re.match(line)
        if m_b:
            indent_spaces = len(m_b.group(1))
            style_name = "List Bullet" if indent_spaces < 2 else "List Bullet 2"
            try:
                para = doc.add_paragraph(style=style_name)
            except KeyError:
                para = doc.add_paragraph()
            _add_runs(para, m_b.group(3).strip())
            continue
        para = doc.add_paragraph()
        _add_runs(para, line)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def render_export_button(content: str, key_prefix: str) -> None:
    """アシスタント応答を docx に変換するダウンロードボタン."""
    if not content or not content.strip():
        return
    try:
        docx_bytes = convert_markdown_to_docx(content)
    except Exception as e:  # noqa: BLE001
        st.caption(f"⚠ docx 生成に失敗: {e}")
        return
    st.download_button(
        label="📥 この回答を Word (.docx) で保存",
        data=docx_bytes,
        file_name=f"audit_{key_prefix}.docx",
        mime=DOCX_MIME,
        key=f"docx_{key_prefix}",
    )


def build_attachment_from_upload(uploaded_file) -> dict:
    """Streamlit UploadedFile → Anthropic content block."""
    name = uploaded_file.name
    suffix = Path(name).suffix.lower().lstrip(".")
    data = uploaded_file.getvalue()
    b64 = base64.standard_b64encode(data).decode("utf-8")
    if suffix == "pdf":
        return {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": b64,
            },
        }
    if suffix in IMAGE_MEDIA:
        return {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": IMAGE_MEDIA[suffix],
                "data": b64,
            },
        }
    raise ValueError(
        f"unsupported file type: .{suffix} (PDF/PNG/JPG/GIF/WebP のみ対応)"
    )


def build_user_message(text: str, attachments: list[dict]) -> dict:
    content: list[dict] = list(attachments)
    if text:
        content.append({"type": "text", "text": text})
    if not content:
        raise ValueError("text か attachment のどちらかが必要")
    return {"role": "user", "content": content}


def stream_assistant(
    *,
    model: str,
    effort: str,
    thinking_display: str,
    max_tokens: int,
    messages: list[dict],
    system_blocks: list[dict],
    api_key: str | None = None,
) -> Iterator[tuple[str, object]]:
    """API を呼び出し ('thinking'|'text', str) または ('usage', usage) を yield."""
    client = anthropic.Anthropic(api_key=api_key) if api_key else anthropic.Anthropic()
    with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        thinking={"type": "adaptive", "display": thinking_display},
        output_config={"effort": effort},
        system=system_blocks,
        messages=messages,
        cache_control={"type": "ephemeral"},
    ) as stream:
        for event in stream:
            if event.type == "content_block_delta":
                if event.delta.type == "thinking_delta":
                    yield ("thinking", event.delta.thinking)
                elif event.delta.type == "text_delta":
                    yield ("text", event.delta.text)
        final = stream.get_final_message()
    yield ("usage", final.usage)
    yield ("stop_reason", final.stop_reason)


def summarize_user_msg(msg: dict) -> str:
    parts = []
    for block in msg["content"]:
        t = block.get("type")
        if t == "text":
            parts.append(block["text"])
        elif t == "document":
            kb = len(block["source"]["data"]) // 1024
            parts.append(f"📄 PDF添付 ({kb}KB)")
        elif t == "image":
            kb = len(block["source"]["data"]) // 1024
            mt = block["source"]["media_type"]
            parts.append(f"🖼 画像添付 ({mt}, {kb}KB)")
    return "\n\n".join(parts)


def save_session_to_disk(name: str, session: dict) -> Path:
    p = SESSIONS_DIR / f"{name}.json"
    session["updated_at"] = datetime.now(timezone.utc).isoformat()
    p.write_text(
        json.dumps(session, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return p


def list_saved_sessions() -> list[str]:
    return sorted(p.stem for p in SESSIONS_DIR.glob("*.json"))


@st.cache_data(show_spinner=False)
def index_sessions() -> list[dict]:
    """sessions/ の全 .json を読み、本文テキストでインデックス化."""
    index: list[dict] = []
    for p in sorted(SESSIONS_DIR.glob("*.json")):
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
        texts: list[str] = []
        for m in data.get("messages", []):
            c = m.get("content")
            if isinstance(c, str):
                texts.append(c)
            elif isinstance(c, list):
                for block in c:
                    if isinstance(block, dict) and block.get("type") == "text":
                        texts.append(block.get("text", ""))
        body = "\n".join(texts)
        index.append({
            "name": p.stem,
            "path": str(p),
            "updated_at": data.get("updated_at") or "",
            "created_at": data.get("created_at") or "",
            "turn_count": len(data.get("messages", [])) // 2,
            "body": body,
        })
    index.sort(key=lambda x: x["updated_at"] or x["created_at"], reverse=True)
    return index


def find_snippet(body: str, query: str, span: int = 60) -> str:
    """body 中で query が最初に現れた箇所の前後を返す."""
    if not query:
        return ""
    idx = body.lower().find(query.lower())
    if idx < 0:
        return ""
    start = max(0, idx - span)
    end = min(len(body), idx + len(query) + span)
    snippet = body[start:end].replace("\n", " ").strip()
    prefix = "..." if start > 0 else ""
    suffix = "..." if end < len(body) else ""
    return f"{prefix}{snippet}{suffix}"


# ─────────────────────────────────────────────────────────────────────
#  UI
# ─────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="再下請負通知書 審査アシスタント",
    page_icon="📑",
    layout="wide",
)


# ─── ログインゲート ─────────────────────────────────────────────
AUTH_CONFIG_PATH = ROOT / "auth_config.yaml"


@st.cache_resource(show_spinner=False)
def load_authenticator() -> stauth.Authenticate:
    if not AUTH_CONFIG_PATH.exists():
        st.error(
            f"認証設定ファイル {AUTH_CONFIG_PATH.name} が見つかりません。"
        )
        st.stop()
    with AUTH_CONFIG_PATH.open(encoding="utf-8") as f:
        cfg = yaml.safe_load(f)
    cookie_cfg = cfg.get("cookie") or {}
    cookie_key = cookie_cfg.get("key", "")
    # Streamlit Cloud では secrets.toml の auth.cookie_key で上書き
    secrets_key = ""
    try:
        secrets_key = st.secrets.get("auth", {}).get("cookie_key", "")
    except (FileNotFoundError, AttributeError):
        secrets_key = ""
    if secrets_key:
        cookie_key = secrets_key
    return stauth.Authenticate(
        cfg["credentials"],
        cookie_cfg.get("name", "kainos_audit_auth"),
        cookie_key or "dev-only-please-replace",
        cookie_cfg.get("expiry_days", 7),
    )


authenticator = load_authenticator()
authenticator.login(location="main", fields={
    "Form name": "🔐 ログイン",
    "Username": "ユーザー名",
    "Password": "パスワード",
    "Login": "ログイン",
})

auth_status = st.session_state.get("authentication_status")
if auth_status is False:
    st.error("ユーザー名またはパスワードが正しくありません。")
    st.stop()
elif auth_status is None:
    st.info("ユーザー名とパスワードを入力してください。")
    st.caption(
        "アクセス権がない場合は管理者にお問い合わせください。"
    )
    st.stop()

# 以降は認証済みユーザーのみ実行
_current_user = st.session_state.get("name", "ユーザー")
_current_username = st.session_state.get("username", "")

if "messages" not in st.session_state:
    st.session_state.messages = []
if "session_name" not in st.session_state:
    st.session_state.session_name = "default"
if "last_usage" not in st.session_state:
    st.session_state.last_usage = None
if "api_key" not in st.session_state:
    st.session_state.api_key = ""


def get_effective_api_key() -> tuple[str, str]:
    """APIキーとそのソース名を返す。

    優先順位:
        1. st.secrets[anthropic][api_key]（サーバー側に埋め込まれた共通キー）
        2. st.session_state[api_key]（ユーザーがブラウザから入力）
        3. 環境変数 ANTHROPIC_API_KEY（ローカル開発用）
    """
    try:
        secret_key = st.secrets.get("anthropic", {}).get("api_key", "")
    except (FileNotFoundError, AttributeError):
        secret_key = ""
    if secret_key:
        return secret_key, "サーバー設定"
    if st.session_state.get("api_key"):
        return st.session_state["api_key"], "セッション入力"
    env_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if env_key:
        return env_key, "環境変数"
    return "", ""


def is_server_managed_api_key() -> bool:
    """サーバー側 secrets にキーが設定されているか."""
    try:
        return bool(st.secrets.get("anthropic", {}).get("api_key", ""))
    except (FileNotFoundError, AttributeError):
        return False


with st.sidebar:
    st.markdown(f"### 👤 {_current_user}")
    st.caption(f"username: `{_current_username}`")
    authenticator.logout(button_name="🚪 ログアウト", location="sidebar")
    st.divider()

    effective_key, key_source = get_effective_api_key()
    if is_server_managed_api_key():
        # 管理者がサーバー側に埋め込んでいるためユーザーは入力不要
        st.caption("✅ APIキーはサーバー管理")
    else:
        # ローカル開発 or 個別キー運用
        st.markdown("### 🔑 Anthropic API キー")
        session_key = st.session_state.get("api_key", "")
        st.text_input(
            "API キー",
            type="password",
            key="api_key",
            placeholder="sk-ant-...",
            help=(
                "console.anthropic.com で発行。ブラウザのセッション中のみ保持され、"
                "リロードで揮発します。"
            ),
            label_visibility="collapsed",
        )
        effective_key, key_source = get_effective_api_key()
        if effective_key:
            masked = (
                f"{effective_key[:8]}...{effective_key[-4:]}"
                if len(effective_key) > 14
                else "***"
            )
            st.success(f"✅ 設定済み ({key_source}: `{masked}`)")
        else:
            st.warning("⚠ APIキー未設定")
        if session_key:
            if st.button(
                "🗑 入力キーをクリア",
                key="clear_api_key",
                use_container_width=True,
            ):
                st.session_state.api_key = ""
                st.rerun()

    st.divider()
    st.markdown("### ⚙ 設定")
    model = st.selectbox(
        "モデル",
        ["claude-sonnet-4-6", "claude-opus-4-7", "claude-opus-4-6"],
        index=0,
        help="Sonnet 4.6: 1審査 ~20-30円 / Opus 4.7: ~40-450円（高難度向け）",
    )
    effort = st.selectbox(
        "effort",
        ["low", "medium", "high", "xhigh", "max"],
        index=2,
        help="思考の深さ / トークン消費の上限. high が intelligence と コストの balance",
    )
    max_tokens = st.slider("max_tokens", 1024, 64000, 16000, step=1024)
    thinking_display = st.selectbox(
        "思考過程の表示",
        ["summarized", "omitted"],
        index=0,
    )

    st.divider()
    st.markdown("### 💾 セッション")
    session_name = st.text_input(
        "セッション名",
        value=st.session_state.session_name,
        help="保存ファイル名 (sessions/<name>.json)",
    )
    st.session_state.session_name = session_name
    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("💾 保存", use_container_width=True):
            session_data = {
                "model": model,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": None,
                "messages": st.session_state.messages,
            }
            saved = save_session_to_disk(session_name, session_data)
            index_sessions.clear()
            st.success(f"保存しました: {saved.name}")
    with col_b:
        if st.button("🗑 履歴クリア", use_container_width=True):
            st.session_state.messages = []
            st.session_state.last_usage = None
            st.rerun()

    st.markdown("#### 🔍 過去セッション検索")
    search_q = st.text_input(
        "キーワード",
        placeholder="例: 主任技術者 / 健康保険",
        label_visibility="collapsed",
    )
    sess_index = index_sessions()
    if not sess_index:
        st.caption("保存済みセッションはありません。")
    else:
        if search_q:
            matches = [
                e for e in sess_index
                if search_q.lower() in e["body"].lower()
                or search_q.lower() in e["name"].lower()
            ]
        else:
            matches = sess_index
        st.caption(
            f"全 {len(sess_index)}件中 {len(matches)}件表示 (新しい順)"
        )
        for e in matches[:15]:
            with st.container(border=True):
                head = f"**{e['name']}** ({e['turn_count']} ターン)"
                if e["updated_at"]:
                    head += f"\n\n*{e['updated_at'][:19]}*"
                st.markdown(head)
                if search_q:
                    snip = find_snippet(e["body"], search_q)
                    if snip:
                        st.caption(snip)
                if st.button(
                    "📂 開く",
                    key=f"open_sess_{e['name']}",
                    use_container_width=True,
                ):
                    p = Path(e["path"])
                    data = json.loads(p.read_text(encoding="utf-8"))
                    st.session_state.messages = data.get("messages", [])
                    st.session_state.session_name = e["name"]
                    st.success(f"読み込みました: {e['name']}")
                    st.rerun()

    st.divider()
    st.markdown("### 📚 法令ソース")
    if SOURCES_DIR.exists():
        for pdf in sorted(SOURCES_DIR.glob("*.pdf")):
            kb = pdf.stat().st_size // 1024
            st.caption(f"📄 {pdf.name} ({kb:,}KB)")
    else:
        st.caption("sources/ が見つかりません")

    if st.session_state.last_usage:
        st.divider()
        u = st.session_state.last_usage
        st.markdown("### 📊 直近 usage")
        st.caption(
            f"in: {u.input_tokens:,} / cache_w: "
            f"{u.cache_creation_input_tokens:,} / cache_r: "
            f"{u.cache_read_input_tokens:,} / out: {u.output_tokens:,}"
        )

# Main panel
st.title("📑 再下請負通知書 審査アシスタント")
st.caption(
    "建設業法 / 国交省ガイドラインのコーパス全文 (~427K tokens) を背景に、"
    "通知書の審査と Q&A を行います。"
)

_eff_key, _eff_src = get_effective_api_key()
_api_key_ready = bool(_eff_key)
if not _api_key_ready:
    if is_server_managed_api_key():
        st.info("🔑 サーバー側 API キーが空欄です。管理者にお問い合わせください。")
    else:
        st.info(
            "🔑 まずサイドバーに **Anthropic API キー** を入力してください "
            "(`sk-ant-...` で始まる文字列)。"
        )


def run_turn(text: str, files) -> None:
    """ユーザー入力を1ターン送信して履歴に追加。失敗時は履歴を変更しない."""
    api_key, _ = get_effective_api_key()
    if not api_key:
        st.error(
            "⚠ APIキー未設定。"
            "管理者にお問い合わせください "
            "(またはサイドバーからキーを入力)。"
        )
        return
    try:
        attachments = [build_attachment_from_upload(f) for f in (files or [])]
        user_msg = build_user_message(text, attachments)
    except ValueError as e:
        st.error(f"入力エラー: {e}")
        return

    candidate = st.session_state.messages + [user_msg]
    with st.chat_message("user"):
        st.markdown(summarize_user_msg(user_msg))

    system_blocks = build_system_blocks()

    with st.chat_message("assistant"):
        thinking_box = (
            st.expander("🤔 思考過程", expanded=False)
            if thinking_display == "summarized"
            else None
        )
        thinking_ph = thinking_box.empty() if thinking_box else None
        response_ph = st.empty()
        usage_ph = st.empty()

        thinking_buf: list[str] = []
        response_buf: list[str] = []
        final_usage = None
        try:
            for kind, val in stream_assistant(
                model=model,
                effort=effort,
                thinking_display=thinking_display,
                max_tokens=max_tokens,
                messages=candidate,
                system_blocks=system_blocks,
                api_key=api_key,
            ):
                if kind == "thinking" and thinking_ph is not None:
                    thinking_buf.append(val)
                    thinking_ph.markdown("".join(thinking_buf))
                elif kind == "text":
                    response_buf.append(val)
                    response_ph.markdown(
                        highlight_citations("".join(response_buf)),
                        unsafe_allow_html=True,
                    )
                elif kind == "usage":
                    final_usage = val
                    usage_ph.caption(
                        f"in: {val.input_tokens:,} / cache_r: "
                        f"{val.cache_read_input_tokens:,} / "
                        f"out: {val.output_tokens:,}"
                    )
        except anthropic.APIError as e:
            st.error(f"API エラー: {e}")
            return

        full_response = "".join(response_buf)
        render_checklist_summary(full_response, position="above")
        cits = extract_citations(full_response)
        live_key = f"live{len(st.session_state.messages)}"
        render_citations(cits, key_prefix=live_key)
        render_export_button(full_response, key_prefix=live_key)

    # 成功時のみ履歴に反映
    st.session_state.messages.append(user_msg)
    st.session_state.messages.append(
        {"role": "assistant", "content": full_response}
    )
    if final_usage is not None:
        st.session_state.last_usage = final_usage


def run_batch_audit(files, *, model, effort, max_tokens) -> list[dict]:
    """各添付に AUDIT_PROMPT を投入し、結果リストを返す."""
    api_key, _ = get_effective_api_key()
    if not api_key:
        st.error("⚠ APIキー未設定。管理者にお問い合わせください。")
        return []
    system_blocks = build_system_blocks()
    results: list[dict] = []
    progress = st.progress(0.0, text="開始準備中...")
    log_area = st.empty()
    total = len(files)
    for i, f in enumerate(files):
        progress.progress(i / total, text=f"審査中 ({i + 1}/{total}): {f.name}")
        log_area.caption(f"📄 {f.name} を送信中...")
        try:
            attachment = build_attachment_from_upload(f)
        except ValueError as e:
            results.append({"name": f.name, "error": str(e)})
            continue
        user_msg = {
            "role": "user",
            "content": [attachment, {"type": "text", "text": AUDIT_PROMPT}],
        }
        try:
            response_buf: list[str] = []
            usage_obj = None
            for kind, val in stream_assistant(
                model=model,
                effort=effort,
                thinking_display="omitted",
                max_tokens=max_tokens,
                messages=[user_msg],
                system_blocks=system_blocks,
                api_key=api_key,
            ):
                if kind == "text":
                    response_buf.append(val)
                elif kind == "usage":
                    usage_obj = val
            full = "".join(response_buf)
            results.append({
                "name": f.name,
                "md": full,
                "sections": parse_audit_sections(full),
                "usage": usage_obj,
            })
        except anthropic.APIError as e:
            results.append({"name": f.name, "error": f"API: {e}"})
    progress.progress(1.0, text=f"完了 ({total}件)")
    log_area.empty()
    return results


# chat_input は root レベル限定（tabs 内には置けない）
user_text = st.chat_input(
    "質問を入力 (例: 主任技術者の専任要件は？)",
    disabled=not _api_key_ready,
)

tab_chat, tab_batch = st.tabs(
    ["💬 対話 / 個別審査", "📦 一括審査"],
)

with tab_chat:
    # 過去履歴を表示
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg["role"] == "user":
                st.markdown(summarize_user_msg(msg))
            else:
                content = msg["content"]
                if isinstance(content, list):
                    content = "".join(
                        b["text"] for b in content if b.get("type") == "text"
                    )
                render_checklist_summary(content, position="history")
                st.markdown(
                    highlight_citations(content),
                    unsafe_allow_html=True,
                )
                cits = extract_citations(content)
                hist_idx = st.session_state.messages.index(msg)
                render_citations(cits, key_prefix=f"hist{hist_idx}")
                render_export_button(content, key_prefix=f"hist{hist_idx}")

    # 入力エリア
    st.divider()
    upload_col, audit_col = st.columns([3, 1])
    with upload_col:
        uploaded_files = st.file_uploader(
            "📎 通知書を添付 (PDF/PNG/JPG/GIF/WebP, 複数可)",
            type=["pdf", "png", "jpg", "jpeg", "gif", "webp"],
            accept_multiple_files=True,
            key="chat_uploader",
        )
    with audit_col:
        st.markdown("&nbsp;")
        run_audit = st.button(
            "🔍 構造化審査を実行",
            use_container_width=True,
            type="primary",
            disabled=(not uploaded_files) or (not _api_key_ready),
            help="添付した書類に対し AUDIT_PROMPT を投入して厳格審査を行います",
        )

    # dispatch
    if run_audit and uploaded_files:
        run_turn(AUDIT_PROMPT, uploaded_files)
    elif user_text:
        run_turn(user_text, uploaded_files)

with tab_batch:
    st.markdown("### 📦 通知書の一括審査")
    st.caption(
        "複数の通知書をまとめてアップロードし、一括で構造化審査を実行します。"
        "システムプロンプト (~427K tokens) は ephemeral cache を共有するため、"
        "2件目以降は cache_read が効きます。"
    )

    if "batch_results" not in st.session_state:
        st.session_state.batch_results = []

    batch_col1, batch_col2 = st.columns([3, 1])
    with batch_col1:
        batch_files = st.file_uploader(
            "通知書をまとめて添付",
            type=["pdf", "png", "jpg", "jpeg", "gif", "webp"],
            accept_multiple_files=True,
            key="batch_uploader",
        )
    with batch_col2:
        st.markdown("&nbsp;")
        run_batch = st.button(
            "🚀 一括審査を実行",
            type="primary",
            use_container_width=True,
            disabled=(not batch_files) or (not _api_key_ready),
        )

    if run_batch and batch_files:
        st.session_state.batch_results = run_batch_audit(
            batch_files,
            model=model,
            effort=effort,
            max_tokens=max_tokens,
        )

    results = st.session_state.batch_results
    if results:
        st.divider()
        col_h, col_clr = st.columns([5, 1])
        with col_h:
            st.markdown(f"### 📊 結果サマリー ({len(results)}件)")
        with col_clr:
            if st.button("🗑 結果クリア", key="batch_clear"):
                st.session_state.batch_results = []
                st.rerun()

        rows = []
        for r in results:
            if "error" in r:
                rows.append({
                    "ファイル": r["name"],
                    "状態": f"⚠ {r['error'][:50]}",
                    "🚨重大": "-",
                    "⚠記載漏れ": "-",
                    "✅正常": "-",
                    "🔍要再確認": "-",
                })
            else:
                s = r["sections"]
                rows.append({
                    "ファイル": r["name"],
                    "状態": "✓ 完了",
                    "🚨重大": len(s.get("critical", [])),
                    "⚠記載漏れ": len(s.get("warning", [])),
                    "✅正常": len(s.get("ok", [])),
                    "🔍要再確認": len(s.get("unclear", [])),
                })
        st.dataframe(rows, hide_index=True, use_container_width=True)

        st.divider()
        st.markdown("### 📄 各書類の詳細")
        for i, r in enumerate(results):
            if "error" in r:
                with st.expander(f"⚠ {r['name']} (エラー)", expanded=False):
                    st.error(r["error"])
                continue
            s = r["sections"]
            badge = (
                f"🚨{len(s.get('critical', []))} "
                f"⚠{len(s.get('warning', []))} "
                f"✅{len(s.get('ok', []))}"
            )
            with st.expander(
                f"📄 {r['name']}  —  {badge}",
                expanded=False,
            ):
                render_checklist_summary(r["md"], position="above")
                st.markdown(
                    highlight_citations(r["md"]),
                    unsafe_allow_html=True,
                )
                cits = extract_citations(r["md"])
                render_citations(cits, key_prefix=f"batch{i}")
                render_export_button(r["md"], key_prefix=f"batch{i}")
                if r.get("usage"):
                    u = r["usage"]
                    st.caption(
                        f"usage — in: {u.input_tokens:,} / "
                        f"cache_r: {u.cache_read_input_tokens:,} / "
                        f"out: {u.output_tokens:,}"
                    )
